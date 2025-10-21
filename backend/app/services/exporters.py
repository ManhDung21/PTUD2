"""Utilities to export generated descriptions into various formats."""

from datetime import datetime
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def export_docx(text: str) -> BytesIO:
    """Return an in-memory DOCX file containing the description."""
    doc = Document()

    title = doc.add_heading('Mô Tả Sản Phẩm', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    timestamp = doc.add_paragraph(f'Ngày tạo: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_pdf(text: str) -> BytesIO:
    """Return an in-memory PDF file containing the description."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()

    title = Paragraph("Mô Tả Sản Phẩm", styles['Heading1'])
    story.append(title)
    story.append(Spacer(1, 0.2 * inch))

    timestamp = Paragraph(f'Ngày tạo: {datetime.now().strftime("%d/%m/%Y %H:%M")}', styles['Normal'])
    story.append(timestamp)
    story.append(Spacer(1, 0.3 * inch))

    for line in text.split('\n'):
        if line.strip():
            story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 0.1 * inch))

    doc.build(story)
    buffer.seek(0)
    return buffer
